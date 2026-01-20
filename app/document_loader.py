"""
Document Loader Module
Handles loading and chunking of PDF, Word, and Excel documents.
"""
import os
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


def load_pdf(file_path: str) -> List[Document]:
    """Load text content from a PDF file."""
    from pypdf import PdfReader
    
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    
    return [Document(
        page_content=text,
        metadata={"source": os.path.basename(file_path), "type": "pdf"}
    )]


def load_docx(file_path: str) -> List[Document]:
    """Load text content from a Word document."""
    from docx import Document as DocxDocument
    
    doc = DocxDocument(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    return [Document(
        page_content=text,
        metadata={"source": os.path.basename(file_path), "type": "docx"}
    )]


def load_excel(file_path: str) -> List[Document]:
    """Load text content from an Excel file."""
    from openpyxl import load_workbook
    
    workbook = load_workbook(file_path, data_only=True)
    text_parts = []
    
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        text_parts.append(f"Sheet: {sheet_name}")
        
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
            if row_text.strip():
                text_parts.append(row_text)
    
    return [Document(
        page_content="\n".join(text_parts),
        metadata={"source": os.path.basename(file_path), "type": "excel"}
    )]


def load_document(file_path: str) -> List[Document]:
    """
    Auto-detect file type and load document.
    Returns a list of Document objects.
    """
    extension = os.path.splitext(file_path)[1].lower()
    
    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".xlsx": load_excel,
        ".xls": load_excel,
    }
    
    loader = loaders.get(extension)
    if not loader:
        raise ValueError(f"Unsupported file type: {extension}")
    
    return loader(file_path)


def chunk_documents(documents: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """
    Split documents into smaller chunks for embedding.
    
    Args:
        documents: List of Document objects
        chunk_size: Maximum size of each chunk
        chunk_overlap: Overlap between consecutive chunks
    
    Returns:
        List of chunked Document objects
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    return text_splitter.split_documents(documents)
